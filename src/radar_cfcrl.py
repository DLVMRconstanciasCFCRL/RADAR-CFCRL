from playwright.sync_api import sync_playwright
from rapidfuzz import fuzz
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import unicodedata
import hashlib
import os
import json
import smtplib
from email.message import EmailMessage

def limpiar_texto_correo(valor):
    if valor is None:
        return ""

    valor = str(valor)
    valor = valor.replace("\xa0", " ")
    valor = valor.replace("\u200b", "")
    valor = valor.strip()

    return valor


def limpiar_password_correo(valor):
    if valor is None:
        return ""

    valor = str(valor)
    valor = valor.replace("\xa0", "")
    valor = valor.replace(" ", "")
    valor = valor.replace("\u200b", "")
    valor = valor.strip()

    return valor


try:
    import streamlit as st

    CORREO_REMITENTE = limpiar_texto_correo(st.secrets["CORREO_REMITENTE"])
    PASSWORD_CORREO = limpiar_password_correo(st.secrets["PASSWORD_CORREO"])
    CORREO_DESTINATARIO = limpiar_texto_correo(st.secrets["CORREO_DESTINATARIO"])

except Exception:
    CORREO_REMITENTE = limpiar_texto_correo(os.getenv("CORREO_REMITENTE", ""))
    PASSWORD_CORREO = limpiar_password_correo(os.getenv("PASSWORD_CORREO", ""))
    CORREO_DESTINATARIO = limpiar_texto_correo(os.getenv("CORREO_DESTINATARIO", ""))

URL = "https://servicios.centrolaboral.gob.mx/constancia-de-representatividad/solicitudes"

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CLIENTES_PATH = os.path.join(BASE_DIR, "clientes", "clientes.csv")
EVIDENCIAS_DIR = os.path.join(BASE_DIR, "evidencias")
REPORTES_DIR = os.path.join(BASE_DIR, "reportes")

os.makedirs(EVIDENCIAS_DIR, exist_ok=True)
os.makedirs(REPORTES_DIR, exist_ok=True)


def normalizar(texto):
    if texto is None:
        return ""

    texto = str(texto).lower()
    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(c for c in texto if not unicodedata.combining(c))

    texto = texto.replace("s.a. de c.v.", "sa de cv")
    texto = texto.replace("s. de r.l. de c.v.", "s de rl de cv")
    texto = texto.replace("s de r.l. de c.v.", "s de rl de cv")
    texto = texto.replace("wal-mart", "wal mart")
    texto = texto.replace("walmart", "wal mart")

    caracteres_a_quitar = [".", ",", ";", ":", "(", ")", "\n", "\t", "/", "\\", "_"]

    for caracter in caracteres_a_quitar:
        texto = texto.replace(caracter, " ")

    texto = " ".join(texto.split())
    return texto


def cargar_clientes():
    clientes = pd.read_csv(CLIENTES_PATH)
    lista_clientes = []

    for _, row in clientes.iterrows():
        keywords = str(row["keywords"]).split(";")

        lista_clientes.append({
            "cliente": row["cliente"],
            "keywords": [k.strip() for k in keywords if k.strip()]
        })

    return lista_clientes


def detectar_coincidencias(texto, clientes):
    texto_normalizado = normalizar(texto)
    hallazgos = []

    for cliente in clientes:
        for keyword in cliente["keywords"]:
            keyword_normalizada = normalizar(keyword)

            if keyword_normalizada in texto_normalizado:
                hallazgos.append({
                    "cliente": cliente["cliente"],
                    "keyword": keyword,
                    "tipo": "coincidencia exacta",
                    "score": 100
                })
                continue

            if len(keyword_normalizada) >= 18:
                score = fuzz.partial_ratio(keyword_normalizada, texto_normalizado)

                if score >= 96:
                    hallazgos.append({
                        "cliente": cliente["cliente"],
                        "keyword": keyword,
                        "tipo": "coincidencia aproximada",
                        "score": score
                    })

    return hallazgos
    
def avanzar_pagina(page):
    pudo_avanzar = page.evaluate("""
        () => {
            const elementos = Array.from(document.querySelectorAll('a, button, span'));

            const candidatos = elementos.filter(el => {
                const texto = (el.textContent || '').trim();
                const aria = (el.getAttribute('aria-label') || '').trim().toLowerCase();
                const title = (el.getAttribute('title') || '').trim().toLowerCase();

                return texto === '»'
                    || texto === '›'
                    || texto === 'Siguiente'
                    || texto === 'siguiente'
                    || aria.includes('siguiente')
                    || aria.includes('next')
                    || title.includes('siguiente')
                    || title.includes('next');
            });

            if (candidatos.length === 0) {
                return false;
            }

            const boton = candidatos[candidatos.length - 1];

            const claseBoton = boton.className ? String(boton.className).toLowerCase() : '';
            const clasePadre = boton.parentElement && boton.parentElement.className
                ? String(boton.parentElement.className).toLowerCase()
                : '';

            const estaDeshabilitado =
                boton.disabled
                || claseBoton.includes('disabled')
                || clasePadre.includes('disabled')
                || boton.getAttribute('aria-disabled') === 'true';

            if (estaDeshabilitado) {
                return false;
            }

            boton.click();
            return true;
        }
    """)

    return pudo_avanzar


def crear_reporte_word(reporte):
    fecha_archivo = datetime.now().strftime("%Y%m%d_%H%M%S")
    word_path = os.path.join(REPORTES_DIR, f"reporte_cfcrl_{fecha_archivo}.docx")

    document = Document()

    titulo = document.add_heading("Reporte de Monitoreo CFCRL", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitulo = document.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run("Solicitudes de Constancia de Representatividad")
    run.bold = True
    run.font.size = Pt(12)

    document.add_paragraph("")

    p = document.add_paragraph()
    p.add_run("Fecha de consulta: ").bold = True
    p.add_run(reporte["fecha_consulta"])

    p = document.add_paragraph()
    p.add_run("Fuente revisada: ").bold = True
    p.add_run(reporte["url"])

    p = document.add_paragraph()
    p.add_run("Total de hallazgos: ").bold = True
    p.add_run(str(reporte["total_hallazgos"]))

    document.add_paragraph("")

    if reporte["total_hallazgos"] == 0:
        document.add_heading("Resultado", level=1)
        document.add_paragraph(
            "No se detectaron coincidencias relevantes con los clientes registrados en la base interna."
        )
    else:
        document.add_heading("Alertas detectadas", level=1)

        tabla = document.add_table(rows=1, cols=5)
        tabla.style = "Table Grid"

        encabezados = tabla.rows[0].cells
        encabezados[0].text = "Cliente"
        encabezados[1].text = "Palabra detectada"
        encabezados[2].text = "Tipo de coincidencia"
        encabezados[3].text = "Score"
        encabezados[4].text = "Página CFCRL"

        clientes_agregados = set()

        for hallazgo in reporte["hallazgos"]:
            pagina_cfcrl = str(hallazgo.get("pagina_cfcrl", "No identificada"))

            for coincidencia in hallazgo["coincidencias"]:
                clave = (
                    coincidencia["cliente"],
                    coincidencia["keyword"],
                    coincidencia["tipo"],
                    str(coincidencia["score"]),
                    pagina_cfcrl
                )

                if clave in clientes_agregados:
                    continue

                clientes_agregados.add(clave)

                fila = tabla.add_row().cells
                fila[0].text = coincidencia["cliente"]
                fila[1].text = coincidencia["keyword"]
                fila[2].text = coincidencia["tipo"]
                fila[3].text = str(coincidencia["score"])
                fila[4].text = pagina_cfcrl

        document.add_paragraph("")

        document.add_heading("Acción sugerida", level=1)
        document.add_paragraph(
            "Revisar manualmente la publicación en la página del Centro Federal de Conciliación y Registro Laboral, "
            "confirmar la empresa, sindicato solicitante, fecha de publicación y domicilio del centro de trabajo. "
            "El presente reporte es únicamente una alerta de monitoreo y no debe considerarse como notificación procesal."
        )

    document.add_paragraph("")
    document.add_heading("Evidencia generada", level=1)

    p = document.add_paragraph()
    p.add_run("Captura de pantalla: ").bold = True
    p.add_run(reporte["captura"])

    p = document.add_paragraph()
    p.add_run("Archivo HTML: ").bold = True
    p.add_run(reporte["html"])

    document.save(word_path)

    return word_path


def enviar_correo_reporte(reporte, word_path):
    asunto = "Reporte CFCRL sin hallazgos"

    if reporte["total_hallazgos"] > 0:
        asunto = "ALERTA CFCRL: posibles coincidencias detectadas"

    cuerpo = f"""
Estimado Oscar,

Se ejecutó el monitoreo automático de avisos de solicitud de constancia de representatividad del Centro Federal de Conciliación y Registro Laboral.

Fecha de consulta: {reporte["fecha_consulta"]}

Fuente revisada:
{reporte["url"]}

Total de hallazgos detectados: {reporte["total_hallazgos"]}

Se adjunta el reporte en formato Word, así como las rutas internas de evidencia generadas por el sistema.

Este correo constituye únicamente una alerta de monitoreo interno y no debe considerarse como notificación procesal.

Saludos.
Radar CFCRL por Oscar de la Vega Castillo
"""

    mensaje = EmailMessage()
    mensaje["From"] = CORREO_REMITENTE
    mensaje["To"] = CORREO_DESTINATARIO
    mensaje["Subject"] = asunto
    mensaje.set_content(cuerpo)

    with open(word_path, "rb") as archivo:
        contenido = archivo.read()
        nombre_archivo = os.path.basename(word_path)

    mensaje.add_attachment(
        contenido,
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=nombre_archivo
    )

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
        smtp.login(CORREO_REMITENTE, PASSWORD_CORREO)
        smtp.send_message(mensaje)


def revisar_cfcrl():
    clientes = cargar_clientes()

    fecha = datetime.now()
    fecha_archivo = fecha.strftime("%Y%m%d_%H%M%S")

    screenshot_path = os.path.join(EVIDENCIAS_DIR, f"cfcrl_{fecha_archivo}.png")
    html_path = os.path.join(EVIDENCIAS_DIR, f"cfcrl_{fecha_archivo}.html")
    reporte_path = os.path.join(REPORTES_DIR, f"reporte_{fecha_archivo}.json")

    resultados = []
    paginas_revisadas = []

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()

        page.goto(URL, wait_until="domcontentloaded", timeout=60000)
        page.wait_for_timeout(8000)

        numero_pagina = 1

        while True:
            print(f"Revisando página {numero_pagina}")

            page.wait_for_timeout(2500)

            texto_actual = page.inner_text("body")

            paginas_revisadas.append({
                "pagina": numero_pagina,
                "texto": texto_actual
            })

            if numero_pagina >= 100:
                print("Se llegó al límite de 100 páginas.")
                break

            pudo_avanzar = avanzar_pagina(page)

            if not pudo_avanzar:
                print("No se encontró botón de siguiente o ya no hay más páginas.")
                break

            numero_pagina += 1
            page.wait_for_timeout(3500)

        page.screenshot(path=screenshot_path, full_page=True)

        html = page.content()

        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html)

        hash_html = hashlib.sha256(html.encode("utf-8")).hexdigest()

        for pagina_revisada in paginas_revisadas:
            numero_pagina_detectada = pagina_revisada["pagina"]
            texto_pagina_individual = pagina_revisada["texto"]

            texto_completo = " ".join(
                [linea.strip() for linea in texto_pagina_individual.split("\n") if linea.strip()]
            )

            coincidencias_pagina = detectar_coincidencias(texto_completo, clientes)

            if coincidencias_pagina:
                resultados.append({
                    "fecha_consulta": fecha.isoformat(),
                    "url": URL,
                    "pagina_cfcrl": numero_pagina_detectada,
                    "texto_detectado": texto_completo,
                    "coincidencias": coincidencias_pagina,
                    "captura": screenshot_path,
                    "html": html_path,
                    "hash_html": hash_html
                })

        browser.close()

    reporte = {
        "fecha_consulta": fecha.isoformat(),
        "url": URL,
        "total_hallazgos": len(resultados),
        "hallazgos": resultados,
        "captura": screenshot_path,
        "html": html_path
    }

    with open(reporte_path, "w", encoding="utf-8") as f:
        json.dump(reporte, f, ensure_ascii=False, indent=4)

    return reporte


if __name__ == "__main__":
    reporte = revisar_cfcrl()
    word_path = crear_reporte_word(reporte)

    try:
        enviar_correo_reporte(reporte, word_path)
        correo_enviado = True
        error_correo = ""
    except Exception as e:
        correo_enviado = False
        error_correo = str(e)

    print("")
    print("REPORTE CFCRL")
    print("Fecha de consulta:", reporte["fecha_consulta"])
    print("Total de hallazgos:", reporte["total_hallazgos"])

    if reporte["total_hallazgos"] > 0:
        print("")
        print("ALERTAS DETECTADAS")

        for hallazgo in reporte["hallazgos"]:
            print("")
            print("Página CFCRL:", hallazgo.get("pagina_cfcrl", "No identificada"))
            print("Coincidencias:")

            for coincidencia in hallazgo["coincidencias"]:
                print(
                    coincidencia["cliente"],
                    "|",
                    coincidencia["keyword"],
                    "|",
                    coincidencia["tipo"],
                    "|",
                    coincidencia["score"]
                )

    else:
        print("")
        print("Sin coincidencias relevantes.")

    print("")
    print("Reporte Word generado en:")
    print(word_path)

    print("")

    if correo_enviado:
        print("Correo enviado a:")
        print(CORREO_DESTINATARIO)
    else:
        print("No se pudo enviar el correo.")
        print("Error:")
        print(error_correo)
        
